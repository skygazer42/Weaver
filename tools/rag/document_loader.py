"""
Document Loader for RAG Pipeline.

Parses various document formats into text chunks for embedding.
"""

import hashlib
import logging
import os
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

logger = logging.getLogger(__name__)

# Check for optional dependencies
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    fitz = None
    PYMUPDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DocxDocument = None
    DOCX_AVAILABLE = False


@dataclass
class Document:
    """A document chunk with metadata."""
    content: str
    metadata: Dict[str, Any] = field(default_factory=dict)
    doc_id: str = ""
    chunk_id: str = ""

    def __post_init__(self):
        if not self.doc_id:
            self.doc_id = hashlib.md5(self.content[:500].encode()).hexdigest()[:12]
        if not self.chunk_id:
            self.chunk_id = f"{self.doc_id}_{hashlib.md5(self.content.encode()).hexdigest()[:8]}"


class DocumentLoader:
    """
    Load and parse documents from various formats.

    Supports:
    - PDF (via PyMuPDF)
    - DOCX (via python-docx)
    - TXT, MD (plain text)
    """

    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
    ):
        """
        Initialize the document loader.

        Args:
            chunk_size: Maximum characters per chunk
            chunk_overlap: Overlap between chunks for context continuity
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def load(self, file_path: Union[str, Path]) -> List[Document]:
        """
        Load a document from file path.

        Args:
            file_path: Path to the document

        Returns:
            List of Document chunks
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        suffix = path.suffix.lower()
        text = ""
        metadata = {
            "source": str(path),
            "filename": path.name,
            "file_type": suffix,
        }

        if suffix == ".pdf":
            text = self._load_pdf(path)
        elif suffix in (".docx", ".doc"):
            text = self._load_docx(path)
        elif suffix in (".txt", ".md", ".markdown", ".rst"):
            text = self._load_text(path)
        else:
            # Try as plain text
            try:
                text = self._load_text(path)
            except Exception as e:
                raise ValueError(f"Unsupported file format: {suffix}") from e

        if not text.strip():
            logger.warning(f"No text extracted from {path}")
            return []

        return self._chunk_text(text, metadata)

    def load_from_bytes(
        self,
        content: bytes,
        filename: str,
        file_type: str = None,
    ) -> List[Document]:
        """
        Load a document from bytes.

        Args:
            content: File content as bytes
            filename: Original filename
            file_type: File type (pdf, docx, txt, etc.)

        Returns:
            List of Document chunks
        """
        if file_type is None:
            file_type = Path(filename).suffix.lower()

        metadata = {
            "source": filename,
            "filename": filename,
            "file_type": file_type,
        }

        text = ""

        if file_type in (".pdf", "pdf"):
            text = self._parse_pdf_bytes(content)
        elif file_type in (".docx", ".doc", "docx", "doc"):
            text = self._parse_docx_bytes(content)
        else:
            try:
                text = content.decode("utf-8")
            except UnicodeDecodeError:
                text = content.decode("latin-1")

        if not text.strip():
            logger.warning(f"No text extracted from {filename}")
            return []

        return self._chunk_text(text, metadata)

    def _load_pdf(self, path: Path) -> str:
        """Load text from PDF file."""
        if not PYMUPDF_AVAILABLE:
            raise ImportError(
                "PyMuPDF is required for PDF parsing. "
                "Install with: pip install pymupdf"
            )

        text_parts = []
        with fitz.open(path) as doc:
            for page_num, page in enumerate(doc, 1):
                page_text = page.get_text()
                if page_text.strip():
                    text_parts.append(f"[Page {page_num}]\n{page_text}")

        return "\n\n".join(text_parts)

    def _parse_pdf_bytes(self, content: bytes) -> str:
        """Parse PDF from bytes."""
        if not PYMUPDF_AVAILABLE:
            raise ImportError("PyMuPDF is required for PDF parsing.")

        text_parts = []
        with fitz.open(stream=content, filetype="pdf") as doc:
            for page_num, page in enumerate(doc, 1):
                page_text = page.get_text()
                if page_text.strip():
                    text_parts.append(f"[Page {page_num}]\n{page_text}")

        return "\n\n".join(text_parts)

    def _load_docx(self, path: Path) -> str:
        """Load text from DOCX file."""
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for DOCX parsing. "
                "Install with: pip install python-docx"
            )

        doc = DocxDocument(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)

    def _parse_docx_bytes(self, content: bytes) -> str:
        """Parse DOCX from bytes."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX parsing.")

        import io
        doc = DocxDocument(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)

    def _load_text(self, path: Path) -> str:
        """Load plain text file."""
        try:
            return path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            return path.read_text(encoding="latin-1")

    def _chunk_text(
        self,
        text: str,
        metadata: Dict[str, Any],
    ) -> List[Document]:
        """
        Split text into chunks with overlap.

        Uses sentence-aware splitting when possible.
        """
        # Clean up text
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = re.sub(r" {2,}", " ", text)

        chunks = []
        doc_id = hashlib.md5(text[:1000].encode()).hexdigest()[:12]

        if len(text) <= self.chunk_size:
            return [Document(
                content=text.strip(),
                metadata=metadata,
                doc_id=doc_id,
            )]

        # Split by paragraphs first
        paragraphs = text.split("\n\n")
        current_chunk = ""

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            if len(current_chunk) + len(para) + 2 <= self.chunk_size:
                current_chunk = f"{current_chunk}\n\n{para}" if current_chunk else para
            else:
                if current_chunk:
                    chunks.append(Document(
                        content=current_chunk.strip(),
                        metadata={**metadata, "chunk_index": len(chunks)},
                        doc_id=doc_id,
                    ))
                    # Overlap: keep last part of current chunk
                    overlap_text = current_chunk[-self.chunk_overlap:] if len(current_chunk) > self.chunk_overlap else ""
                    current_chunk = f"{overlap_text}\n\n{para}" if overlap_text else para
                else:
                    # Paragraph itself is too long, split by sentences
                    sentences = re.split(r"(?<=[.!?])\s+", para)
                    for sent in sentences:
                        if len(current_chunk) + len(sent) + 1 <= self.chunk_size:
                            current_chunk = f"{current_chunk} {sent}" if current_chunk else sent
                        else:
                            if current_chunk:
                                chunks.append(Document(
                                    content=current_chunk.strip(),
                                    metadata={**metadata, "chunk_index": len(chunks)},
                                    doc_id=doc_id,
                                ))
                            current_chunk = sent

        # Don't forget the last chunk
        if current_chunk.strip():
            chunks.append(Document(
                content=current_chunk.strip(),
                metadata={**metadata, "chunk_index": len(chunks)},
                doc_id=doc_id,
            ))

        logger.info(f"Split document into {len(chunks)} chunks")
        return chunks


def load_documents(
    paths: List[Union[str, Path]],
    chunk_size: int = 1000,
    chunk_overlap: int = 200,
) -> List[Document]:
    """
    Load multiple documents.

    Args:
        paths: List of file paths
        chunk_size: Maximum characters per chunk
        chunk_overlap: Overlap between chunks

    Returns:
        List of all Document chunks
    """
    loader = DocumentLoader(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    all_docs = []

    for path in paths:
        try:
            docs = loader.load(path)
            all_docs.extend(docs)
            logger.info(f"Loaded {len(docs)} chunks from {path}")
        except Exception as e:
            logger.error(f"Failed to load {path}: {e}")

    return all_docs
